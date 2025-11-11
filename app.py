import streamlit as st
import pandas as pd
import docx
from docx.shared import Pt
import random
import time
import unicodedata

@st.cache_data
def cargar_libros():
    try:
        df = pd.read_csv("libros_titulo.csv", header=None)
        df.columns = ["T√≠tulo"]
        return df["T√≠tulo"].tolist()
    except Exception as e:
        st.error("Error al cargar el archivo libros_titulo.csv.")
        return []

def analizar_libro(titulo):
    time.sleep(2)
    return {
        "palabras": random.randint(500, 1500),
        "estructuras": ["oraciones simples", "coordinadas", "subordinadas"],
        "complejidad": random.choice(["Baja", "Media", "Alta"]),
        "vocabulario": random.choice(["Com√∫n", "Con modismos"]),
        "tiempos": random.choice(["Presente", "Pasado", "Futuro"]),
        "nivel_dificultad": random.choice(["Inicial", "Intermedio", "Avanzado"])
    }

def generar_ficha(titulo, analisis):
    doc = docx.Document()
    doc.add_heading(f"Ficha de comprensi√≥n lectora: {titulo}", 0)
    doc.add_paragraph(f"N√∫mero de palabras: {analisis['palabras']}")
    doc.add_paragraph(f"Estructuras gramaticales: {', '.join(analisis['estructuras'])}")
    doc.add_paragraph(f"Complejidad del texto: {analisis['complejidad']}")
    doc.add_paragraph(f"Vocabulario: {analisis['vocabulario']}")
    doc.add_paragraph(f"Tiempos verbales predominantes: {analisis['tiempos']}")
    doc.add_paragraph(f"Nivel de dificultad: {analisis['nivel_dificultad']}")
    doc.add_paragraph("1. ¬øQu√© sucede al inicio del libro?")
    doc.add_paragraph("2. ¬øQui√©nes son los personajes principales?")
    doc.add_paragraph("3. ¬øQu√© ense√±anza deja el libro?")
    doc.save("ficha.docx")
    return "ficha.docx"

def generar_rubrica(titulo):
    doc = docx.Document()
    doc.add_heading(f"R√∫brica de evaluaci√≥n: {titulo}", 0)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Criterio"
    hdr_cells[1].text = "Excelente"
    hdr_cells[2].text = "Bueno"
    hdr_cells[3].text = "Necesita mejorar"
    criterios = ["Comprensi√≥n global", "Identificaci√≥n de personajes", "Secuencia de eventos", "Reflexi√≥n personal"]
    for criterio in criterios:
        row_cells = table.add_row().cells
        row_cells[0].text = criterio
        row_cells[1].text = "Responde con profundidad y claridad"
        row_cells[2].text = "Responde con cierta claridad"
        row_cells[3].text = "Respuestas incompletas o confusas"
    doc.save("rubrica.docx")
    return "rubrica.docx"

def recomendar_libros(titulo, lista_libros):
    recomendados = random.sample([l for l in lista_libros if l != titulo], min(5, len(lista_libros)-1))
    recomendaciones = []
    for libro in recomendados:
        recomendaciones.append({
            "t√≠tulo": libro,
            "pros": ["Favorece vocabulario", "Estimula la imaginaci√≥n", "Adecuado para el nivel lector"],
            "contras": ["Puede requerir apoyo adulto", "Vocabulario avanzado"]
        })
    return recomendaciones

# Funci√≥n para normalizar texto (eliminar acentos y convertir a min√∫sculas)
def normalizar_texto(texto):
    texto = texto.lower()
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join([c for c in texto if not unicodedata.combining(c)])
    return texto

st.set_page_config(page_title="An√°lisis pedag√≥gico y Recomendaciones", layout="centered")
st.title("üìö An√°lisis pedag√≥gico y Recomendaciones")

libros_disponibles = cargar_libros()
titulo_libro = st.text_input("Introduce el t√≠tulo del libro infantil")
buscar = st.button("üîç Analizar")

if buscar and titulo_libro:
    # Normalizar lista de libros y t√≠tulo introducido
    libros_normalizados = [normalizar_texto(l) for l in libros_disponibles]
    titulo_libro_normalizado = normalizar_texto(titulo_libro)

    if titulo_libro_normalizado not in libros_normalizados:
        st.warning("El libro no se encuentra en la base de datos.")
    else:
        indice = libros_normalizados.index(titulo_libro_normalizado)
        titulo_original = libros_disponibles[indice]

        with st.spinner("Analizando el libro..."):
            analisis = analizar_libro(titulo_original)

            st.subheader("üîé An√°lisis pedag√≥gico")
            st.write(f"**N√∫mero de palabras:** {analisis['palabras']}")
            st.write(f"**Estructuras gramaticales:** {', '.join(analisis['estructuras'])}")
            st.write(f"**Complejidad del texto:** {analisis['complejidad']}")
            st.write(f"**Vocabulario:** {analisis['vocabulario']}")
            st.write(f"**Tiempos verbales predominantes:** {analisis['tiempos']}")
            st.write(f"**Nivel de dificultad:** {analisis['nivel_dificultad']}")

            st.subheader("üìÑ Ficha de comprensi√≥n lectora")
            ficha_path = generar_ficha(titulo_original, analisis)
            with open(ficha_path, "rb") as f:
                st.download_button("üì• Descargar ficha", f, file_name=ficha_path)

            st.subheader("üìä R√∫brica de evaluaci√≥n")
            rubrica_path = generar_rubrica(titulo_original)
            with open(rubrica_path, "rb") as f:
                st.download_button("üì• Descargar r√∫brica", f, file_name=rubrica_path)

            st.subheader("üìö Recomendaciones pedag√≥gicas")
            recomendaciones = recomendar_libros(titulo_original, libros_disponibles)
            for rec in recomendaciones:
                st.markdown(f"**T√≠tulo:** {rec['t√≠tulo']}")
                st.markdown(f"‚úÖ Pros: {', '.join(rec['pros'])}")
                st.markdown(f"‚ö†Ô∏è Contras: {', '.join(rec['contras'])}")
                st.markdown("---")
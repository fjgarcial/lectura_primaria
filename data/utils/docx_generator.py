from docx import Document

def generar_ficha_comprension(titulo, contenido, preguntas):
    doc = Document()
    doc.add_heading(f'Ficha de Comprensión: {titulo}', 0)
    doc.add_paragraph(contenido)
    doc.add_heading('Preguntas de Comprensión', level=1)
    for i, pregunta in enumerate(preguntas, 1):
        doc.add_paragraph(f"{i}. {pregunta}")
    nombre_archivo = f"ficha_{titulo.replace(' ', '_')}.docx"
    doc.save(nombre_archivo)
    return nombre_archivo
